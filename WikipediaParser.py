# -*- coding: utf-8 -*-
import sys, click, os
import wikipediaapi as wikiapi
import docx

__author__ = 'QuatreB'
__version__ = 'v1.0'

defFont = 'Times New Roman'
defSize = docx.shared.Pt(14)
defHeadingSize = docx.shared.Pt(24)
defFirstLineOffset = docx.shared.Cm(1.25)
defLineSpacing = 1

wikiLang = 'en'

missDict = {
    'ar':
        ('مصادر'),
    'bg':
        ('Външни препратки'),
    'id':
        ('Referensi'),
    'de':
        ('Weblinks', 'Einzelnachweise', 'Enlaces externos'),
    'fr':
        ('Voir aussi', 'Références'),
    'es':
        ('Véase también', 'Notas y referencias'),
    'nt':
        ('Externe links'),
    'pl':
        ('Zobacz też', 'Linki zewnętrzne', 'Przypisy'),
    'en':
        ('See also', 'References', 'External links'),
    'ru':
        ('См. также', 'Примечания', 'Ссылки'),
    'uk':
        ('Примітки', 'Посилання')
    }

def AddParagraphs(doc, text):
    for paragraph in text.split('\n'):
        doc.add_paragraph(paragraph)

def AddSections(doc, sections, level = 1):
    for section in sections:
        if wikiLang not in missDict:
            doc.add_heading(section.title, level)
            AddParagraphs(doc, section.text)
            AddSections(doc, section.sections, level + 1)
        elif section.title not in missDict[wikiLang]:
            p = doc.add_heading(section.title, level)
            #print(f'   -- Added heading {section.title}')
            AddParagraphs(doc, section.text)
            AddSections(doc, section.sections, level + 1)

def CreateDocument(page):
    doc = docx.Document()
    for style in doc.styles:
        if style.type == docx.enum.style.WD_STYLE_TYPE.PARAGRAPH:
            style.font.name = defFont
            style.font.size = defSize
            style.paragraph_format.first_line_indent = defFirstLineOffset
            style.paragraph_format.line_spacing = defLineSpacing
            if style.name.startswith('Heading'):
                style.paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                style.font.size = defHeadingSize - docx.shared.Pt(2 * (int(style.name[-1])))
            elif style.name == 'Title':
                style.font.size = defHeadingSize
                style.font.bold = True
                style.paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_heading(page.title, 0)
    AddParagraphs(doc, page.summary)
    AddSections(doc, page.sections)
    doc.save(page.language + '-' + page.title + '.docx')

@click.command()
@click.argument('page_name', type = str)
@click.option('--lang', '-l',
              default = wikiLang,
              help = 'Wikipedia language (default - \'en\')',
              type = str)
@click.option('--font', '-f',
              default = defFont,
              help = 'Default font for Word document (default - \'Times New Roman\')',
              type = str)

def main(page_name, lang, font):
    '''
    This scipt parse Wikipedia page in Word document
    '''
    try:
        global wikiLang, defFont
        wikiLang = lang
        defFont = font

        wiki = wikiapi.Wikipedia(
            language = 'en' if wikiLang == 'all' else wikiLang,
            extract_format = wikiapi.ExtractFormat.WIKI
            )

        page = wiki.page(page_name)
        fileName = page.title if wikiLang == 'en' or wikiLang == 'all' else page.langlinks['en'].title
        if not page.exists(): raise Exception('Page doesn\'t exist.')
        try:
            os.chdir(fileName)
        except:
            os.mkdir(fileName)
            os.chdir(fileName)
        if wikiLang == 'all':
            with click.progressbar(iterable = sorted(page.langlinks.items()),
                                   label = 'Creating documents:',
                                   empty_char = ' '
                                   ) as bar:
                for key, value in bar:
                    #print(f'{value.language}: {value.title}')
                    #continue
                    if not value.exists(): raise Exception('PageLink doesn\'t exist.')
                    wikiLang = value.language
                    CreateDocument(value)
        else:
            wikiLang = 'en'
            print(f'  - Creating document {page.language}-{page.title}')
            CreateDocument(page)
    except Exception as e:
        print(f' Error - {e}')

if __name__ == '__main__':
    click.secho(f'\nWikipedia Parser {__version__} \u2014 {__author__} \u00A9 Copyright', fg = 'cyan')
    main()
